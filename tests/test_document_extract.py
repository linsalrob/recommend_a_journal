from __future__ import annotations

from pathlib import Path

import pytest

from journal_recommender.document_extract import extract_manuscript

REPO_ROOT = Path(__file__).resolve().parents[1]
FIXTURES = REPO_ROOT / "tests/fixtures/manuscripts"


def test_txt_extraction_detects_title_abstract_and_sections() -> None:
    extracted = extract_manuscript(FIXTURES / "microbiome_cohort.txt")

    assert extracted.file_type == "txt"
    assert extracted.title.startswith("Synthetic gut microbiome cohort")
    assert "metagenomes" in extracted.abstract
    assert "Methods" in extracted.sections
    assert "Data Availability" in extracted.sections
    assert extracted.warnings == []


def test_md_extraction_detects_sections() -> None:
    extracted = extract_manuscript(FIXTURES / "environmental_metagenomics.md")

    assert extracted.file_type == "md"
    assert "environmental survey" in extracted.abstract.lower()
    assert "Methods" in extracted.sections
    assert "Data Availability" in extracted.sections


def test_docx_extraction_detects_title_and_abstract(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "manuscript.docx"
    document = Document()
    document.add_paragraph("Synthetic DOCX manuscript on phage genomics")
    document.add_paragraph("Alice Example, Bob Example")
    document.add_paragraph("Abstract")
    document.add_paragraph("This synthetic DOCX manuscript studies bacteriophages.")
    document.add_paragraph("Methods")
    document.add_paragraph("We used host prediction and comparative genomics.")
    document.save(path)

    extracted = extract_manuscript(path)

    assert extracted.file_type == "docx"
    assert extracted.title.startswith("Synthetic DOCX manuscript on phage genomics")
    assert "bacteriophages" in extracted.abstract
    assert "Methods" in extracted.sections
    assert not extracted.warnings


def test_pdf_extraction_warns_on_empty_text(tmp_path: Path) -> None:
    pytest.importorskip("pypdf")
    from pypdf import PdfWriter

    path = tmp_path / "empty.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as handle:
        writer.write(handle)

    extracted = extract_manuscript(path)

    assert extracted.file_type == "pdf"
    assert extracted.full_text.strip() == ""
    assert any("PDF text extraction" in warning for warning in extracted.warnings)

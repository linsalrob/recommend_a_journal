"""Local manuscript document extraction helpers."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

HEADING_NAMES = {
    "abstract",
    "introduction",
    "background",
    "methods",
    "materials and methods",
    "results",
    "discussion",
    "conclusion",
    "conclusions",
    "data availability",
    "code availability",
    "software availability",
    "references",
    "supplementary information",
}

TITLE_SKIP_PREFIXES = (
    "author",
    "authors",
    "corresponding author",
    "running title",
    "keywords",
)


@dataclass
class ExtractedManuscript:
    filename: str
    file_type: str
    title: str
    abstract: str
    sections: dict[str, str] = field(default_factory=dict)
    full_text: str = ""
    warnings: list[str] = field(default_factory=list)


def extract_manuscript(
    path: str | Path, raw_bytes: bytes | None = None
) -> ExtractedManuscript:
    path = Path(path)
    file_type = path.suffix.lower().lstrip(".")
    if file_type in {"txt", "md"}:
        text = (
            raw_bytes.decode("utf-8", errors="ignore")
            if raw_bytes is not None
            else path.read_text(encoding="utf-8", errors="ignore")
        )
        return extract_from_text(text, path.name, file_type)
    if file_type == "docx":
        return extract_from_docx(path, raw_bytes=raw_bytes)
    if file_type == "pdf":
        return extract_from_pdf(path, raw_bytes=raw_bytes)
    raise ValueError(f"Unsupported manuscript file type: {path.suffix}")


def extract_from_text(text: str, filename: str, file_type: str) -> ExtractedManuscript:
    lines = normalise_lines(text)
    sections = split_sections(lines)
    title = infer_title(lines, sections)
    abstract = sections.get("Abstract", "").strip()
    warnings: list[str] = []
    if not abstract:
        warnings.append("Abstract not detected; feature extraction may be incomplete.")
    return ExtractedManuscript(
        filename=filename,
        file_type=file_type,
        title=title,
        abstract=abstract,
        sections=sections,
        full_text="\n".join(lines).strip(),
        warnings=warnings,
    )


def extract_from_docx(
    path: Path,
    raw_bytes: bytes | None = None,
) -> ExtractedManuscript:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "python-docx is required to extract DOCX manuscripts."
        ) from exc

    document = Document(path) if raw_bytes is None else Document(_bytes_io(raw_bytes))
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = clean_line(paragraph.text)
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "")
        if is_heading_style(style_name, text):
            paragraphs.append(f"\n{canonical_heading(text)}\n")
        else:
            paragraphs.append(text)
    text = "\n".join(paragraphs)
    extracted = extract_from_text(text, path.name, "docx")
    if len(extracted.full_text) < 500:
        extracted.warnings.append("DOCX text extraction produced a short text body.")
    return extracted


def extract_from_pdf(
    path: Path,
    raw_bytes: bytes | None = None,
) -> ExtractedManuscript:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - import guard
        raise RuntimeError("pypdf is required to extract PDF manuscripts.") from exc

    reader = PdfReader(_bytes_io(raw_bytes) if raw_bytes is not None else str(path))
    pages: list[str] = []
    warnings: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        pages.append(page_text)
    text = "\n".join(pages)
    if len(text.strip()) < 200:
        warnings.append("PDF text extraction may be incomplete.")
    if not text.strip():
        warnings.append("PDF text extraction failed; upload DOCX or paste text.")
    extracted = extract_from_text(text, path.name, "pdf")
    extracted.warnings.extend(warnings)
    return extracted


def split_sections(lines: list[str]) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current = "Preamble"
    sections[current] = []
    for line in lines:
        canonical = canonical_heading(line)
        if canonical:
            current = canonical
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return {
        name: "\n".join(content).strip()
        for name, content in sections.items()
        if "\n".join(content).strip()
    }


def infer_title(lines: list[str], sections: dict[str, str]) -> str:
    abstract_index = first_heading_index(lines, "Abstract")
    candidate_lines = lines[:abstract_index] if abstract_index is not None else lines
    for line in candidate_lines:
        if not line or is_noisy_metadata_line(line):
            continue
        if looks_like_title(line):
            return line
    for line in candidate_lines:
        if line and not is_author_line(line):
            return line
    return ""


def first_heading_index(lines: list[str], heading: str) -> int | None:
    target = heading.lower()
    for index, line in enumerate(lines):
        if canonical_heading(line).lower() == target:
            return index
    return None


def normalise_lines(text: str) -> list[str]:
    return [clean_line(line) for line in text.splitlines()]


def clean_line(line: str) -> str:
    text = " ".join(line.strip().split())
    return strip_markdown_heading(text)


def strip_markdown_heading(text: str) -> str:
    return re.sub(r"^#+\s*", "", text)


def canonical_heading(text: str) -> str:
    stripped = strip_markdown_heading(clean_line(text)).rstrip(":")
    if not stripped:
        return ""
    lower = stripped.lower()
    if lower in HEADING_NAMES:
        return stripped.title() if lower != "rna" else stripped
    if is_heading_like(stripped):
        return stripped
    return ""


def is_heading_style(style_name: str, text: str) -> bool:
    style = style_name.lower()
    return "heading" in style or (text and len(text) < 80 and text.isupper())


def is_heading_like(text: str) -> bool:
    lower = text.lower().rstrip(":")
    return lower in HEADING_NAMES


def looks_like_title(line: str) -> bool:
    if len(line) < 10 or len(line) > 180:
        return False
    lower = line.lower()
    if any(lower.startswith(prefix) for prefix in TITLE_SKIP_PREFIXES):
        return False
    return bool(sum(char.isalpha() for char in line) >= 5)


def is_author_line(line: str) -> bool:
    lower = line.lower()
    return (
        any(
            token in lower
            for token in ["@", "corresponding author", "department", "university", ","]
        )
        and len(line) < 160
    )


def is_noisy_metadata_line(line: str) -> bool:
    lower = line.lower()
    return any(token in lower for token in ["author information", "copyright", "doi"])


def _bytes_io(raw_bytes: bytes | None) -> Any:
    from io import BytesIO

    return BytesIO(raw_bytes or b"")
